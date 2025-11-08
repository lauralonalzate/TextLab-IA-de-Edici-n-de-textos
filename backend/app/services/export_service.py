"""
Export Service for generating PDF and DOCX documents.

Handles document export with institutional templates and APA 7 formatting.
"""
import os
import uuid
import logging
from pathlib import Path
from typing import Dict, Any, Optional, List
from datetime import datetime

from app.core.storage import get_storage

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available. DOCX export will not work.")

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch, pt
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.enums import TA_LEFT, TA_JUSTIFY
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    logging.warning("reportlab not available. PDF export will not work.")

try:
    import textstat
    TEXTSTAT_AVAILABLE = True
except ImportError:
    TEXTSTAT_AVAILABLE = False
    logging.warning("textstat not available. Readability metrics will not be calculated.")

logger = logging.getLogger(__name__)


class ExportService:
    """Service for exporting documents to PDF and DOCX"""
    
    def __init__(self, storage_path: str = "exports"):
        from app.core.storage import LocalStorage
        self.storage = get_storage()
        # Keep storage_path for backward compatibility
        self.storage_path = Path(storage_path)
        # If using LocalStorage, ensure directory exists
        if isinstance(self.storage, LocalStorage):
            self.storage_path.mkdir(parents=True, exist_ok=True)
    
    def export_to_docx(
        self,
        document: Dict[str, Any],
        references: List[Dict[str, Any]],
        options: Dict[str, Any],
        filename: str
    ) -> str:
        """
        Export document to DOCX format with APA 7 styling.
        
        Args:
            document: Document data dictionary
            references: List of reference dictionaries
            options: Export options (include_stats, template_id, etc.)
            filename: Filename for the exported file
        
        Returns:
            Filename of generated file
        """
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx is not available")
        
        doc = Document()
        
        # Set default font (Inter/Roboto fallback)
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'  # Fallback if Inter/Roboto not available
        font.size = Pt(12)
        
        # Add header with logo (if template specified)
        if options.get("template_id"):
            self._add_header(doc, options.get("template_id"))
        
        # Add title
        title_para = doc.add_heading(document.get("title", "Untitled Document"), level=1)
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Add content
        content = document.get("content", "")
        if content:
            # Split content into paragraphs
            paragraphs = content.split('\n\n')
            for para_text in paragraphs:
                if para_text.strip():
                    para = doc.add_paragraph(para_text.strip())
                    para_format = para.paragraph_format
                    para_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                    para_format.space_after = Pt(6)
        
        # Add references section
        if references:
            doc.add_page_break()
            doc.add_heading('References', level=1)
            
            # Add each reference with hanging indent
            for ref in references:
                ref_text = self._format_reference_for_docx(ref)
                ref_para = doc.add_paragraph(ref_text)
                
                # Apply hanging indent (APA 7 style)
                para_format = ref_para.paragraph_format
                para_format.first_line_indent = Inches(-0.5)  # -18pt ≈ -0.25 inches
                para_format.left_indent = Inches(0.5)  # 36pt ≈ 0.5 inches
                para_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                para_format.space_after = Pt(6)
        
        # Add metadata/stats if requested
        if options.get("include_stats", False):
            doc.add_page_break()
            doc.add_heading('Document Statistics', level=1)
            stats = self._calculate_stats(document.get("content", ""))
            for key, value in stats.items():
                doc.add_paragraph(f"{key}: {value}")
        
        # Save document to temporary location first
        import tempfile
        import os
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            tmp_path = tmp_file.name
            doc.save(tmp_path)
            
            # Read file content
            with open(tmp_path, 'rb') as f:
                docx_bytes = f.read()
            
            # Save using storage backend
            self.storage.save(filename, docx_bytes)
            
            # Clean up temp file
            os.unlink(tmp_path)
        
        return filename
    
    def export_to_pdf(
        self,
        document: Dict[str, Any],
        references: List[Dict[str, Any]],
        options: Dict[str, Any],
        filename: str
    ) -> str:
        """
        Export document to PDF format with APA 7 styling.
        
        Args:
            document: Document data dictionary
            references: List of reference dictionaries
            options: Export options
            filename: Filename for the exported file
        
        Returns:
            Filename of generated file
        """
        if not REPORTLAB_AVAILABLE:
            raise RuntimeError("reportlab is not available")
        
        # Create PDF document (will use temp file)
        import tempfile
        import os
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
            tmp_path = tmp_file.name
        
        pdf = SimpleDocTemplate(
            tmp_path,
            pagesize=letter,
            rightMargin=inch,
            leftMargin=inch,
            topMargin=inch,
            bottomMargin=inch
        )
        
        # Build content
        story = []
        
        # Add title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=getSampleStyleSheet()['Heading1'],
            fontSize=14,
            spaceAfter=12,
            alignment=TA_LEFT
        )
        story.append(Paragraph(document.get("title", "Untitled Document"), title_style))
        story.append(Spacer(1, 12))
        
        # Add content
        content = document.get("content", "")
        if content:
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=getSampleStyleSheet()['Normal'],
                fontSize=12,
                leading=24,  # Double spacing (12pt font, 24pt line height)
                spaceAfter=6,
                alignment=TA_JUSTIFY
            )
            paragraphs = content.split('\n\n')
            for para_text in paragraphs:
                if para_text.strip():
                    story.append(Paragraph(para_text.strip(), normal_style))
                    story.append(Spacer(1, 6))
        
        # Add references
        if references:
            story.append(PageBreak())
            story.append(Paragraph('References', getSampleStyleSheet()['Heading1']))
            story.append(Spacer(1, 12))
            
            for ref in references:
                ref_text = self._format_reference_for_pdf(ref)
                # Create style with hanging indent
                ref_style = ParagraphStyle(
                    'Reference',
                    parent=getSampleStyleSheet()['Normal'],
                    fontSize=12,
                    leading=24,
                    leftIndent=36,  # 36pt left indent
                    firstLineIndent=-18,  # -18pt first line (hanging indent)
                    spaceAfter=6,
                    alignment=TA_LEFT
                )
                story.append(Paragraph(ref_text, ref_style))
                story.append(Spacer(1, 6))
        
        # Add stats if requested
        if options.get("include_stats", False):
            story.append(PageBreak())
            story.append(Paragraph('Document Statistics', getSampleStyleSheet()['Heading1']))
            story.append(Spacer(1, 12))
            stats = self._calculate_stats(document.get("content", ""))
            for key, value in stats.items():
                story.append(Paragraph(f"{key}: {value}", getSampleStyleSheet()['Normal']))
                story.append(Spacer(1, 6))
        
        # Build PDF
        pdf.build(story)
        
        # Read file content
        with open(tmp_path, 'rb') as f:
            pdf_bytes = f.read()
        
        # Save using storage backend
        self.storage.save(filename, pdf_bytes)
        
        # Clean up temp file
        os.unlink(tmp_path)
        
        return filename
    
    def _format_reference_for_docx(self, ref: Dict[str, Any]) -> str:
        """Format reference for DOCX output"""
        from app.services.apa_service import apa7_service
        return apa7_service.generate_reference(ref)
    
    def _format_reference_for_pdf(self, ref: Dict[str, Any]) -> str:
        """Format reference for PDF output"""
        from app.services.apa_service import apa7_service
        return apa7_service.generate_reference(ref)
    
    def _add_header(self, doc: Document, template_id: str):
        """Add institutional header with logo"""
        # This would load template-specific header
        # For now, just add a simple header
        section = doc.sections[0]
        header = section.header
        header_para = header.paragraphs[0]
        header_para.text = "TextLab - Academic Writing Platform"
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _calculate_stats(self, content: str) -> Dict[str, Any]:
        """Calculate document statistics"""
        stats = {
            "Word Count": len(content.split()),
            "Character Count": len(content),
            "Paragraph Count": len([p for p in content.split('\n\n') if p.strip()]),
        }
        
        if TEXTSTAT_AVAILABLE and content:
            try:
                stats["Flesch Reading Ease"] = round(textstat.flesch_reading_ease(content), 2)
                stats["Flesch-Kincaid Grade Level"] = round(textstat.flesch_kincaid_grade(content), 2)
                stats["Automated Readability Index"] = round(textstat.automated_readability_index(content), 2)
            except Exception as e:
                logger.warning(f"Error calculating readability metrics: {e}")
        
        return stats
    
    def generate_filename(self, document_id: str, format: str) -> str:
        """Generate unique filename for export"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        return f"document_{document_id}_{timestamp}.{format}"
    
    def get_file_path(self, filename: str) -> Path:
        """Get full path for a filename"""
        return self.storage_path / filename


# Singleton instance
export_service = ExportService()

